import PyPDF2
from docx import Document
import logging
from typing import Optional
import os

logger = logging.getLogger(__name__)

class TextExtractor:
    """Extract text from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'doc', 'txt']
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from file based on its extension
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text or None if extraction fails
        """
        try:
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._extract_from_docx(file_path)
            elif file_extension == 'txt':
                return self._extract_from_txt(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {str(e)}")
                        continue
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise
        
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file with latin-1 encoding: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            raise
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate if file exists and has supported format
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            bool: True if file is valid, False otherwise
        """
        if not os.path.exists(file_path):
            logger.error(f"File does not exist: {file_path}")
            return False
        
        file_extension = file_path.lower().split('.')[-1]
        if file_extension not in self.supported_formats:
            logger.error(f"Unsupported file format: {file_extension}")
            return False
        
        return True